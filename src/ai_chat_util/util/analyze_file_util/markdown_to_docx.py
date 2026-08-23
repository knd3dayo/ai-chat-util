from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

from .markdown_parser import MarkdownParser


class MarkdownToDocxUtil:
    @classmethod
    def convert_markdown_to_docx(cls, markdown_text: str, output_path: str | Path) -> str:
        document = Document()
        document.styles["Normal"].font.name = "Arial"
        document.styles["Normal"].font.size = 11

        blocks = MarkdownParser.parse(markdown_text)
        cls._render_ast_to_docx(document, blocks)

        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output))
        return str(output)

    @classmethod
    def _render_ast_to_docx(cls, document: Document, ast: Any) -> None:
        if isinstance(ast, list):
            for item in ast:
                cls._render_ast_to_docx(document, item)
            return

        if hasattr(ast, "type"):
            node_type = ast.type
            if node_type == "heading":
                level = ast.level
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                run = paragraph.add_run(ast.text)
                run.bold = level <= 2
                run.italic = level == 3
                if level == 1:
                    run.font.size = 20
                elif level == 2:
                    run.font.size = 16
                else:
                    run.font.size = 13
                return

            if node_type == "paragraph":
                if ast.text:
                    document.add_paragraph(ast.text)
                return

            if node_type == "list_item":
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.add_run(ast.text)
                return

            if node_type == "code_block":
                paragraph = document.add_paragraph()
                paragraph.add_run(ast.code)
                return

            if node_type == "table":
                if ast.rows:
                    max_cols = max(len(row) for row in ast.rows)
                    table = document.add_table(rows=1, cols=max_cols)
                    table.style = "Table Grid"
                    for row_index, row in enumerate(ast.rows):
                        if row_index > 0:
                            table.add_row()
                        for col_index, cell in enumerate(row):
                            cell_obj = table.rows[row_index].cells[col_index]
                            cell_obj.text = cell
                            for paragraph in cell_obj.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)
                    for cell in table.rows[0].cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                return

            if node_type == "image":
                if ast.image_path:
                    paragraph = document.add_paragraph()
                    try:
                        image_path = Path(ast.image_path)
                        if not image_path.is_absolute():
                            image_path = Path.cwd() / image_path
                        paragraph.add_run().add_picture(str(image_path), width=Inches(2.5))
                    except FileNotFoundError:
                        paragraph.add_run(f"[image: {ast.image_path}]")
                return

            return
